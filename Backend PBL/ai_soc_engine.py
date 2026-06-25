# ==========================================================
# ai_soc_engine.py
# CypherGuard — AI SOC Engine
# ==========================================================
# Flow:
#   1. Filter SecurityEvent / Alert from DB
#   2. Build structured JSON payload
#   3. Feed  prompt_html_report.txt + JSON  →  Ollama
#   4. Ollama returns a complete HTML page
#   5. Save as  REPORT_TYPE__DATE__DEVICE.html
#   6. Also save .docx (python-docx) from the same AI text
# ==========================================================

import os
import json
import subprocess
import csv
import io
import re
from datetime import datetime, timezone, timedelta
from typing import Optional

from models import db, SecurityEvent, Alert
from logging_engine import log_event

def get_model_name():
    import json
    import os
    config_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "cypherguard.json")
    if os.path.exists(config_path):
        try:
            with open(config_path, "r", encoding="utf-8") as f:
                return json.load(f).get("active_model", "llama3")
        except Exception:
            pass
    return "llama3"

REPORTS_FOLDER = "reports"
PROMPTS_FOLDER = "prompts"

# Single prompt used for ALL report types.
# The JSON payload contains "report_type" so Ollama knows which sections to write.
HTML_PROMPT_FILE = "prompt_html_report.txt"


# ==========================================================
# 1.  Ollama health checks
# ==========================================================

def check_ollama_installed() -> bool:
    try:
        subprocess.run(["ollama", "--version"], check=True, capture_output=True)
        return True
    except (FileNotFoundError, subprocess.CalledProcessError):
        return False


def ensure_model_exists() -> bool:
    try:
        model = get_model_name()
        result = subprocess.run(
            ["ollama", "list"],
            capture_output=True, text=True, check=True, timeout=15
        )
        if model in result.stdout:
            return True

        log_event(event_type="OLLAMA_PULL_START", severity="INFO",
                  message=f"Pulling {model} ...")
        subprocess.run(["ollama", "pull", model], check=True, timeout=600)
        log_event(event_type="OLLAMA_MODEL_PULLED", severity="INFO",
                  message=f"Model {model} ready")
        return True

    except Exception as exc:
        log_event(event_type="OLLAMA_MODEL_ERROR", severity="CRITICAL",
                  message=str(exc))
        return False


# ==========================================================
# 2.  Smart filename builder
#     FULL_SOC__2026-03-16__ESP32_001
#     ALERT_ANALYSIS__2026-03-01_to_2026-03-16__ALL_DEVICES
# ==========================================================

REPORT_LABELS = {
    "full_soc":       "FULL_SOC",
    "alert_analysis": "ALERT_ANALYSIS",
    "risk_summary":   "RISK_SUMMARY",
}


def _sanitize(value: str) -> str:
    return re.sub(r"[^\w\-]", "_", value.strip()).upper()


def _build_basename(report_key: str, filters: dict) -> str:
    os.makedirs(REPORTS_FOLDER, exist_ok=True)

    label    = REPORT_LABELS.get(report_key, report_key.upper())
    specific = filters.get("specific_dates")
    start_s  = filters.get("start_date")
    end_s    = filters.get("end_date")

    if specific:
        date_part = "_".join(specific)
    elif start_s and end_s:
        date_part = "{}_to_{}".format(start_s, end_s)
    elif start_s:
        date_part = start_s
    else:
        date_part = datetime.utcnow().strftime("%Y-%m-%d")

    dev = filters.get("device_id", "")
    device_part = "ALL_DEVICES" if (not dev or dev.lower() in ("", "all", "all devices")) \
                  else _sanitize(dev)

    basename  = "{}__{}__{}".format(label, date_part, device_part)
    base_path = os.path.join(REPORTS_FOLDER, basename)

    if os.path.exists(base_path + ".html") or os.path.exists(base_path + ".docx"):
        ts       = datetime.utcnow().strftime("%H%M%S")
        basename = "{}_{}".format(basename, ts)

    return os.path.join(REPORTS_FOLDER, basename)


# ==========================================================
# 3.  Run Ollama — returns raw output string
# ==========================================================

def _run_ollama(prompt_file: str, json_payload: str) -> str:
    import requests
    import os
    import json
    
    prompt_path = os.path.join(PROMPTS_FOLDER, prompt_file)
    if not os.path.exists(prompt_path):
        raise FileNotFoundError(
            "Prompt file not found: {}".format(prompt_path)
        )

    with open(prompt_path, "r", encoding="utf-8") as fh:
        base_prompt = fh.read()

    full_input = (
        base_prompt
        + "\n\n"
        + "=" * 60 + "\n"
        + "SECURITY DATA (JSON):\n"
        + "=" * 60 + "\n"
        + json_payload
    )

    model = get_model_name()
    
    # Load custom host and threads if configured
    config_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "cypherguard.json")
    ollama_host = "http://127.0.0.1:11434"
    ollama_threads = 2  # Default to 2 to prevent CPU overheating
    if os.path.exists(config_path):
        try:
            with open(config_path, "r", encoding="utf-8") as f:
                cfg = json.load(f)
                ollama_host = cfg.get("ollama_host", "http://127.0.0.1:11434")
                ollama_threads = cfg.get("ollama_threads", 2)
        except Exception:
            pass

    try:
        url = f"{ollama_host}/api/generate"
        headers = {"Content-Type": "application/json"}
        payload = {
            "model": model,
            "prompt": full_input,
            "stream": False,
            "options": {
                "num_thread": ollama_threads
            }
        }
        
        # Generous timeout for local LLM inference
        response = requests.post(url, json=payload, headers=headers, timeout=180)
        if response.status_code != 200:
            raise RuntimeError("Ollama API returned HTTP status {}: {}".format(response.status_code, response.text))
            
        data = response.json()
        output = data.get("response", "")
        if not output:
            raise RuntimeError("Ollama API returned an empty response. Verify model status.")
            
        cleaned_output = output.strip()
        # Clean markdown code block wraps if returned by the model
        cleaned_output = re.sub(r"^```(?:html|xml)?\s*", "", cleaned_output, flags=re.IGNORECASE)
        cleaned_output = re.sub(r"\s*```$", "", cleaned_output)
        return cleaned_output.strip()
        
    except requests.exceptions.RequestException as e:
        raise RuntimeError(
            "Could not connect to Ollama server at {}. "
            "Make sure the service is running ('ollama serve') and reachable. Error: {}".format(ollama_host, str(e))
        )


# ==========================================================
# 4.  Extract plain text from HTML for DOCX fallback
# ==========================================================

def _html_to_plain(html: str) -> str:
    """Strip HTML tags to get plain text for DOCX generation."""
    text = re.sub(r"<style[^>]*>.*?</style>", " ", html, flags=re.DOTALL)
    text = re.sub(r"<script[^>]*>.*?</script>", " ", text, flags=re.DOTALL)
    text = re.sub(r"<[^>]+>", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


# ==========================================================
# 5.  Save HTML report
# ==========================================================

def _save_html(html: str, basename: str) -> str:
    path = basename + ".html"
    with open(path, "w", encoding="utf-8") as fh:
        fh.write(html)
    return path


# ==========================================================
# 6.  Save DOCX report  (python-docx)
# ==========================================================

def _save_docx(
    plain_text: str,
    report_type: str,
    filters: dict,
    summary: Optional[dict],
    basename: str,
) -> str:
    try:
        from docx import Document  # type: ignore[import-untyped]
        from docx.shared import Pt, RGBColor, Inches  # type: ignore[import-untyped]
        from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore[import-untyped]
        from docx.oxml.ns import qn  # type: ignore[import-untyped]
        from docx.oxml import OxmlElement  # type: ignore[import-untyped]

        doc = Document()
        for sec in doc.sections:
            sec.top_margin    = Inches(1)
            sec.bottom_margin = Inches(1)
            sec.left_margin   = Inches(1.2)
            sec.right_margin  = Inches(1.2)

        HDR   = RGBColor(0x1e, 0x3a, 0x5f)
        WHITE = RGBColor(0xff, 0xff, 0xff)
        GREY  = RGBColor(0x94, 0xa3, 0xb8)
        RED   = RGBColor(0xdc, 0x26, 0x26)
        AMBER = RGBColor(0xea, 0x58, 0x0c)

        def shade_cell(cell, hex_color: str) -> None:
            tcPr = cell._tc.get_or_add_tcPr()
            shd  = OxmlElement("w:shd")
            shd.set(qn("w:val"),   "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"),  hex_color)
            tcPr.append(shd)

        type_display = {
            "full_soc":       "Full SOC Report",
            "alert_analysis": "Alert Analysis",
            "risk_summary":   "Risk Summary",
        }.get(report_type, report_type.replace("_", " ").title())

        generated_at = datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC")
        dev          = filters.get("device_id") or "All Devices"
        start_s      = filters.get("start_date") or ""
        end_s        = filters.get("end_date") or ""
        specific     = ", ".join(filters.get("specific_dates") or [])
        date_label   = specific or ("{} to {}".format(start_s, end_s) if start_s else "All Time")
        sev_filter   = filters.get("severity") or "All"

        # Brand
        brand = doc.add_paragraph()
        brand.alignment = WD_ALIGN_PARAGRAPH.CENTER
        br = brand.add_run("CypherGuard  —  AI Security Operations Center")
        br.bold = True
        br.font.size = Pt(9)
        br.font.color.rgb = GREY

        doc.add_paragraph()

        # Title
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tr = title.add_run(type_display.upper())
        tr.bold = True
        tr.font.size = Pt(22)
        tr.font.color.rgb = HDR

        doc.add_paragraph()

        # Metadata table
        tbl = doc.add_table(rows=2, cols=4)
        tbl.style = "Table Grid"
        headers = ["Generated At",  "Device",   "Date Range", "Severity Filter"]
        values  = [generated_at,    dev,        date_label,   sev_filter]
        for i, (h, v) in enumerate(zip(headers, values)):
            hc = tbl.rows[0].cells[i]
            vc = tbl.rows[1].cells[i]
            shade_cell(hc, "1e3a5f")
            hr2 = hc.paragraphs[0].add_run(h)
            hr2.bold = True
            hr2.font.size = Pt(8)
            hr2.font.color.rgb = WHITE
            vr2 = vc.paragraphs[0].add_run(v)
            vr2.font.size = Pt(9)

        doc.add_paragraph()

        # Summary stats (Full SOC only)
        if summary:
            sh = doc.add_paragraph()
            sh.add_run("SUMMARY STATISTICS").font.size = Pt(10)
            sh.runs[0].bold = True
            sh.runs[0].font.color.rgb = HDR

            st = doc.add_table(rows=2, cols=4)
            st.style = "Table Grid"
            slabels = ["Total Events", "Total Alerts", "Critical Alerts", "High Alerts"]
            svals   = [
                str(summary.get("total_events",    "—")),
                str(summary.get("total_alerts",    "—")),
                str(summary.get("critical_alerts", "—")),
                str(summary.get("high_alerts",     "—")),
            ]
            for i, (lbl, val) in enumerate(zip(slabels, svals)):
                lc = st.rows[0].cells[i]
                vc = st.rows[1].cells[i]
                shade_cell(lc, "e2e8f0")
                lc.paragraphs[0].add_run(lbl).font.size = Pt(8)
                lc.paragraphs[0].runs[0].bold = True
                vrun = vc.paragraphs[0].add_run(val)
                vrun.bold = True
                vrun.font.size = Pt(15)
                if lbl == "Critical Alerts":
                    vrun.font.color.rgb = RED
                elif lbl == "High Alerts":
                    vrun.font.color.rgb = AMBER
                else:
                    vrun.font.color.rgb = HDR

            doc.add_paragraph()

        # Body content from plain text
        section_pattern = re.compile(
            r"(?:^|\n)(\d+\.\s+[A-Z][A-Z\s/_\-]+)\n(.*?)(?=\n\d+\.\s+[A-Z]|\Z)",
            re.DOTALL,
        )
        sections = section_pattern.findall(plain_text)

        if sections:
            for heading_text, body_text in sections:
                # Section heading with dark background
                sec_p = doc.add_paragraph()
                sec_r = sec_p.add_run("  {}  ".format(heading_text.strip()))
                sec_r.bold = True
                sec_r.font.size = Pt(10)
                sec_r.font.color.rgb = WHITE
                pPr = sec_p._p.get_or_add_pPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"),   "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"),  "1e3a5f")
                pPr.append(shd)

                for line in body_text.strip().split("\n"):
                    line = line.strip()
                    if not line:
                        doc.add_paragraph()
                        continue
                    if line.startswith(("- ", "* ", "• ")):
                        bp = doc.add_paragraph(style="List Bullet")
                        bp.add_run(line[2:].strip()).font.size = Pt(10)
                    else:
                        cleaned = re.sub(r"\*\*(.+?)\*\*", r"\1", line)
                        p = doc.add_paragraph()
                        p.add_run(cleaned).font.size = Pt(10)

                doc.add_paragraph()
        else:
            for line in plain_text.split("\n"):
                stripped = line.strip()
                if stripped:
                    p = doc.add_paragraph()
                    p.add_run(stripped).font.size = Pt(10)

        # Footer
        doc.add_paragraph()
        fp = doc.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fr = fp.add_run(
            "Generated by CypherGuard AI SOC Engine  ·  Ollama / {}  ·  {}".format(
                get_model_name(), generated_at
            )
        )
        fr.font.size = Pt(8)
        fr.font.color.rgb = GREY
        fr.italic = True

        path = basename + ".docx"
        doc.save(path)
        return path

    except ImportError:
        log_event(event_type="DOCX_SKIP", severity="WARNING",
                  message="python-docx not installed; .docx skipped.")
        return ""
    except Exception as exc:
        log_event(event_type="DOCX_ERROR", severity="WARNING", message=str(exc))
        return ""


# ==========================================================
# 7.  Build CSV
# ==========================================================

def _build_csv(rows: list) -> str:
    if not rows:
        return ""
    out    = io.StringIO()
    writer = csv.DictWriter(out, fieldnames=list(rows[0].keys()))
    writer.writeheader()
    writer.writerows(rows)
    return out.getvalue()


# ==========================================================
# 8.  Query helpers
# ==========================================================

def _date_filter(query, model, filters: dict):
    if filters.get("specific_dates"):
        dates = []
        for ds in filters["specific_dates"]:
            try:
                dates.append(datetime.strptime(ds, "%Y-%m-%d").date())
            except ValueError:
                raise ValueError("Bad date format: {}  (use YYYY-MM-DD)".format(ds))
        return query.filter(db.func.date(model.created_at).in_(dates))

    start_s = filters.get("start_date")
    end_s   = filters.get("end_date")
    if start_s and end_s:
        try:
            start = datetime.strptime(start_s, "%Y-%m-%d").replace(tzinfo=timezone.utc)
            end   = datetime.strptime(end_s, "%Y-%m-%d").replace(tzinfo=timezone.utc) + timedelta(days=1)
        except ValueError:
            raise ValueError("start_date/end_date must be YYYY-MM-DD")
        return query.filter(model.created_at >= start, model.created_at < end)

    return query


def _device_filter_events(query, device_id: str):
    if device_id and device_id.lower() not in ("", "all", "all devices"):
        return query.filter(SecurityEvent.device_id == device_id)
    return query


def _device_filter_alerts(query, device_id: str):
    if device_id and device_id.lower() not in ("", "all", "all devices"):
        return query.filter(Alert.related_device_id == device_id)
    return query


def _severity_filter_events(query, severity: str):
    if severity and severity.lower() not in ("", "all"):
        return query.filter(SecurityEvent.severity == severity.upper())
    return query


def _severity_filter_alerts(query, severity: str):
    if severity and severity.lower() not in ("", "all"):
        return query.filter(Alert.severity == severity.upper())
    return query


# ==========================================================
# 9.  Data builders
# ==========================================================

def _build_full_soc_data(filters: dict):
    q_e = SecurityEvent.query
    q_a = Alert.query

    q_e = _date_filter(q_e, SecurityEvent, filters)
    q_a = _date_filter(q_a, Alert, filters)

    dev = filters.get("device_id", "")
    q_e = _device_filter_events(q_e, dev)
    q_a = _device_filter_alerts(q_a, dev)

    sev = filters.get("severity", "")
    q_e = _severity_filter_events(q_e, sev)
    q_a = _severity_filter_alerts(q_a, sev)

    # Compute correct global stats first
    total_events_count = q_e.count()
    total_alerts_count = q_a.count()
    critical_alerts_count = q_a.filter(Alert.severity == "CRITICAL").count()
    high_alerts_count = q_a.filter(Alert.severity == "HIGH").count()

    # Limit items in prompt payload to keep CPU prompt evaluation fast
    events = q_e.order_by(SecurityEvent.created_at.desc()).limit(30).all()
    alerts = q_a.order_by(Alert.created_at.desc()).limit(15).all()

    event_rows = [
        {
            "event_type": e.event_type,
            "severity":   e.severity,
            "device_id":  e.device_id,
            "ip_address": e.ip_address,
            "message":    e.message,
            "timestamp":  e.created_at.isoformat() if e.created_at else "",
        }
        for e in events
    ]
    alert_rows = [
        {
            "alert_type":  a.alert_type,
            "severity":    a.severity,
            "description": a.description,
            "source_ip":   a.source_ip,
            "device_id":   a.related_device_id,
            "event_count": a.event_count,
            "soar_done":   a.soar_executed,
            "timestamp":   a.created_at.isoformat() if a.created_at else "",
        }
        for a in alerts
    ]

    summary = {
        "total_events":    total_events_count,
        "total_alerts":    total_alerts_count,
        "critical_alerts": critical_alerts_count,
        "high_alerts":     high_alerts_count,
    }

    data = {
        "report_type":     "Full SOC Report",
        "generated_at":    datetime.utcnow().isoformat(),
        "filters":         filters,
        "summary":         summary,
        "security_events": event_rows, # Top 30 recent events for context
        "alerts":          alert_rows, # Top 15 recent alerts for context
    }
    return json.dumps(data, indent=2), event_rows, summary


def _build_alert_data(filters: dict):
    q = Alert.query
    q = _date_filter(q, Alert, filters)
    q = _device_filter_alerts(q, filters.get("device_id", ""))
    q = _severity_filter_alerts(q, filters.get("severity", ""))

    total_alerts = q.count()
    alerts = q.order_by(Alert.created_at.desc()).limit(20).all()
    rows = [
        {
            "alert_type":  a.alert_type,
            "severity":    a.severity,
            "description": a.description,
            "source_ip":   a.source_ip,
            "device_id":   a.related_device_id,
            "event_count": a.event_count,
            "timestamp":   a.created_at.isoformat() if a.created_at else "",
        }
        for a in alerts
    ]
    data = {
        "report_type":  "Alert Analysis",
        "generated_at": datetime.utcnow().isoformat(),
        "filters":      filters,
        "total_alerts": total_alerts,
        "alerts":       rows, # Top 20 alerts
    }
    return json.dumps(data, indent=2), rows


def _build_risk_data(filters: dict):
    total_events = SecurityEvent.query.count()
    total_alerts = Alert.query.count()
    critical_alerts = Alert.query.filter(Alert.severity == "CRITICAL").count()
    high_alerts = Alert.query.filter(Alert.severity == "HIGH").count()
    medium_alerts = Alert.query.filter(Alert.severity == "MEDIUM").count()
    low_alerts = Alert.query.filter(Alert.severity == "LOW").count()
    soar_executed = Alert.query.filter(Alert.soar_executed == True).count()
    
    # Fetch top 25 alerts for breakdown context
    all_alerts = Alert.query.order_by(Alert.created_at.desc()).limit(25).all()
    
    data = {
        "report_type":     "Risk Summary",
        "generated_at":    datetime.utcnow().isoformat(),
        "filters":         filters,
        "total_events":    total_events,
        "total_alerts":    total_alerts,
        "critical_alerts": critical_alerts,
        "high_alerts":     high_alerts,
        "medium_alerts":   medium_alerts,
        "low_alerts":      low_alerts,
        "soar_executed":   soar_executed,
        "alert_breakdown": [
            {"alert_type": a.alert_type, "severity": a.severity}
            for a in all_alerts
        ],
    }
    return json.dumps(data, indent=2)


# ==========================================================
# 10.  PUBLIC API — called by ai_routes.py
# ==========================================================

def generate_full_soc_report(filters: dict):
    """
    Returns: (html_file, docx_file, csv_string, ai_html, summary)
    """
    if not check_ollama_installed():
        raise EnvironmentError(
            "Ollama is not installed. "
            "Download from https://ollama.com/download"
        )
    if not ensure_model_exists():
        raise RuntimeError("Could not prepare model '{}'".format(get_model_name()))

    json_data, event_rows, summary = _build_full_soc_data(filters)
    ai_html  = _run_ollama(HTML_PROMPT_FILE, json_data)
    basename = _build_basename("full_soc", filters)

    html_file  = _save_html(ai_html, basename)
    plain_text = _html_to_plain(ai_html)
    docx_file  = _save_docx(plain_text, "full_soc", filters, summary, basename)
    csv_data   = _build_csv(event_rows)

    log_event(event_type="AI_SOC_REPORT_GENERATED", severity="INFO",
              message="HTML={} | DOCX={} | events={}".format(
                  html_file, docx_file, len(event_rows)))
    return html_file, docx_file, csv_data, ai_html, summary


def generate_alert_analysis(filters: dict):
    """
    Returns: (html_file, docx_file, csv_string, ai_html)
    """
    if not check_ollama_installed():
        raise EnvironmentError("Ollama is not installed.")
    if not ensure_model_exists():
        raise RuntimeError("Could not prepare model '{}'".format(get_model_name()))

    json_data, alert_rows = _build_alert_data(filters)
    ai_html  = _run_ollama(HTML_PROMPT_FILE, json_data)
    basename = _build_basename("alert_analysis", filters)

    html_file  = _save_html(ai_html, basename)
    plain_text = _html_to_plain(ai_html)
    docx_file  = _save_docx(plain_text, "alert_analysis", filters, None, basename)
    csv_data   = _build_csv(alert_rows)

    log_event(event_type="AI_ALERT_ANALYSIS_GENERATED", severity="INFO",
              message="HTML={} | DOCX={} | alerts={}".format(
                  html_file, docx_file, len(alert_rows)))
    return html_file, docx_file, csv_data, ai_html


def generate_risk_summary(filters: dict):
    """
    Returns: (html_file, docx_file, '', ai_html)
    """
    if not check_ollama_installed():
        raise EnvironmentError("Ollama is not installed.")
    if not ensure_model_exists():
        raise RuntimeError("Could not prepare model '{}'".format(get_model_name()))

    json_data = _build_risk_data(filters)
    ai_html   = _run_ollama(HTML_PROMPT_FILE, json_data)
    basename  = _build_basename("risk_summary", filters)

    html_file  = _save_html(ai_html, basename)
    plain_text = _html_to_plain(ai_html)
    docx_file  = _save_docx(plain_text, "risk_summary", filters, None, basename)

    log_event(event_type="AI_RISK_SUMMARY_GENERATED", severity="INFO",
              message="HTML={} | DOCX={}".format(html_file, docx_file))
    return html_file, docx_file, "", ai_html